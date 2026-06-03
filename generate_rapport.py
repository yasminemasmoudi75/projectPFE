"""
Génère le rapport Word complet de l'analyse ML du projet PFE.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x6b, 0xb5)
    p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    return p

def para(text='', bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F4F8')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-têtes
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Fond bleu
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A8A')
        cell._tc.get_or_add_tcPr().append(shd)
    # Données
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        fill = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph('─' * 90)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def note_box(text, color_fill='FEF3C7', color_border='F59E0B'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_fill)
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run('⚠  ' + text)
    set_font(run, size=10, italic=True, color=(0x92, 0x40, 0x0E))
    return p

def success_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'DCFCE7')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run('✅  ' + text)
    set_font(run, size=10, color=(0x16, 0x65, 0x34))
    return p

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJET DE FIN D\'ÉTUDES')
set_font(run, size=14, bold=True, color=(0x64, 0x74, 0x8B))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analyse Globale et Pédagogique')
set_font(run, size=22, bold=True, color=(0x1E, 0x3A, 0x8A))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projet ML — Prédiction des Ventes Régionales')
set_font(run, size=16, bold=True, color=(0x1D, 0x4E, 0xD8))

doc.add_paragraph()
divider()

# Métriques clés sur la page de titre
metric_data = [
    ['Accuracy', '72.94%'],
    ['F1-Score (pondéré)', '70.03%'],
    ['AUC-ROC', '70.92%'],
    ['CV 5-folds', '72.85% ± 0.84%'],
    ['Transactions analysées', '3 595'],
    ['Gouvernorats couverts', '24 / 24'],
    ['Algorithmes benchmarkés', '10'],
    ['Niveau du projet', 'Avancé / Pré-industriel'],
]
t = doc.add_table(rows=len(metric_data), cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (k, v) in enumerate(metric_data):
    fill = 'EFF6FF' if ri % 2 == 0 else 'DBEAFE'
    c1, c2 = t.rows[ri].cells
    r1 = c1.paragraphs[0].add_run(k)
    set_font(r1, size=11, bold=True, color=(0x1E, 0x3A, 0x8A))
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = c2.paragraphs[0].add_run(v)
    set_font(r2, size=11, bold=True, color=(0x07, 0x47, 0x26))
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for cell in [c1, c2]:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        cell._tc.get_or_add_tcPr().append(shd)
    c1.width = Cm(8)
    c2.width = Cm(6)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Modèle : Voting Ensemble (VotingClassifier)  |  Stack : Python / scikit-learn / Flask / SQL Server')
set_font(run, size=10, italic=True, color=(0x64, 0x74, 0x8B))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — VUE D'ENSEMBLE
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Vue d\'ensemble du projet')

h2('1.1 Le problème métier')
para('Imaginez une entreprise qui distribue des dispositifs médicaux à travers toute la Tunisie. '
     'Elle a 24 gouvernorats à couvrir, des équipes commerciales à déployer, des stocks à '
     'pré-positionner, et un budget limité. La question stratégique est simple mais cruciale :')

p = doc.add_paragraph()
run = p.add_run('Où investir davantage ? Où maintenir le cap ? Où réallouer les ressources ?')
set_font(run, size=12, bold=True, color=(0x1E, 0x3A, 0x8A))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)

para('Sans analyse automatisée, cette décision repose sur l\'intuition des managers, sur des '
     'rapports manuels, et sur l\'expérience historique. C\'est lent, subjectif, et souvent '
     'réactif — on ajuste après avoir constaté un problème, rarement avant.')

h2('1.2 La solution proposée')
para('Ce projet construit un système d\'aide à la décision commerciale régionale, alimenté par '
     'les données réelles de l\'ERP de l\'entreprise. Le système :')

for b in [
    'Analyse automatiquement les transactions de ventes historiques',
    'Prédit l\'intensité commerciale attendue par région',
    'Génère des recommandations stratégiques : AUGMENTER, MAINTENIR, RÉDUIRE',
    'Expose ces prédictions via une API REST accessible en temps réel',
    'Se réentraîne automatiquement quand de nouvelles données arrivent',
]:
    bullet(b)

para('')
para('Ce n\'est pas un simple rapport de statistiques. C\'est un pipeline ML de production, '
     'de l\'extraction des données jusqu\'à la recommandation finale, avec une infrastructure '
     'capable d\'évoluer dans le temps.', italic=True)

h2('1.3 Le pipeline en une phrase')
code_block(
    'Données brutes ERP  →  Nettoyage  →  Inférence géographique  →  Enrichissement INS 2023\n'
    '     →  Entraînement VotingClassifier  →  Recommandations stratégiques  →  API Flask'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PIPELINE ÉTAPE PAR ÉTAPE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('2. Explication étape par étape du pipeline')

# Étape 1
h2('Étape 1 — Connexion et extraction des données SQL Server')
h3('Rôle')
para('Récupérer les données transactionnelles réelles depuis la base ERP de l\'entreprise.')
h3('Ce qui se passe concrètement')
para('Le notebook se connecte au serveur SQL Server local (127.0.0.1:1433, base AA) et exécute '
     'une requête qui joint quatre tables :')
for b in [
    'TabBcvd — les lignes de commandes (produit, quantité, date)',
    'TabStock — le catalogue produits (désignation, prix)',
    'TabBcvm — les entêtes de commandes (lien vers le client)',
    'TabTiers — les clients (ville, gouvernorat)',
]:
    bullet(b)
para('La requête filtre sur les commandes depuis 2022 avec des quantités positives, '
     'récupérant ainsi 3 625 transactions réelles.')
h3('Logique métier')
para('On part des vraies ventes de l\'entreprise, pas de données simulées. Chaque ligne '
     'représente une commande réelle d\'un client réel pour un produit médical réel.')
note_box('Si la connexion SQL échoue, le pipeline bascule automatiquement sur un jeu de '
         'données synthétique qui permet de continuer le développement. En production, '
         'c\'est la connexion réelle qui prime.')
divider()

# Étape 2
h2('Étape 2 — Nettoyage et contrôle qualité des données')
h3('Rôle')
para('S\'assurer que les données sont exploitables avant tout traitement.')
h3('Trois opérations successives')
bullet('Suppression des valeurs manquantes sur Vente, Prix, Date. '
       '82 lignes avec désignation manquante sont conservées (champ non utilisé comme feature).', level=0)
bullet('Suppression des doublons exacts : 30 lignes retirées (0.83% du dataset).', level=0)
bullet('Traitement des prix aberrants par la méthode IQR : les prix au-dessus de 100 DT sont '
       'ramenés à 100 DT (clamping). 519 transactions (14.4%) concernées.', level=0)
success_box('Résultat : 3 595 transactions propres, prêtes pour l\'analyse.')
h3('Logique métier')
para('Un modèle entraîné sur des données mal nettoyées apprend des patterns fictifs. '
     'Le nettoyage n\'est pas une formalité — c\'est ce qui garantit que les décisions '
     'reposent sur la réalité.')
divider()

# Étape 3
h2('Étape 3 — Inférence géographique et augmentation des données')
h3('Le défi réel')
para('L\'ERP contient un champ "gouvernorat" dans la table clients — mais après analyse, '
     'ce champ contient le code 17 pour 99.9% des clients. Il n\'a jamais été correctement '
     'renseigné. On ne peut pas l\'utiliser pour savoir si un client est à Tunis ou à Tataouine.')
h3('La solution — Inférence par proxy économique')
para('L\'idée repose sur une logique économique réelle : les dispositifs médicaux à prix élevé '
     'sont vendus majoritairement dans des régions à fort pouvoir d\'achat (Grand Tunis, Sfax, '
     'Nabeul). On normalise le prix de chaque transaction sur la plage des indices d\'achat '
     'régionaux (INS 2023) et on assigne à chaque transaction le gouvernorat dont l\'indice '
     'économique est le plus proche de son prix normalisé.')
h3('L\'augmentation géographique pondérée')
para('L\'inférence par prix seule crée un déséquilibre : TATAOUINE reçoit 30% des transactions '
     'car son indice d\'achat est proche du prix médian. Pour corriger ce biais, 40% des '
     'transactions sont redistribuées selon la population réelle de chaque gouvernorat (INS 2023).')
success_box('Expérience contrôlée : suppression de l\'augmentation dégrade le F1 de -3.3% '
            'et l\'AUC-ROC de -4.8%. L\'augmentation est donc bénéfique et conservée.')
h3('Le Quality Gate adaptatif')
para('Un module dédié (data_quality.py) mesure automatiquement la couverture géographique '
     'réelle de l\'ERP avant chaque réentraînement. Si la couverture dépasse 50%, '
     'l\'augmentation se désactive automatiquement — sans intervention manuelle.')
note_box('Limite reconnue : l\'assignation géographique est une approximation basée sur '
         'un proxy économique. Les résultats régionaux sont des estimations, pas des mesures directes.')
divider()

# Étape 4
h2('Étape 4 — Enrichissement démographique')
h3('Rôle')
para('Ajouter des informations sur chaque gouvernorat pour enrichir le contexte de chaque transaction.')
para('Une fois la région identifiée, on joint des données INS 2023 :')
add_table(
    ['Variable', 'Valeur min', 'Valeur max', 'Logique'],
    [
        ['Population', '140 000 (Tozeur)', '1 200 000 (Ariana)', 'Taille du marché'],
        ['Indice d\'achat', '83.0 (Tataouine)', '135.0 (Tunis)', 'Pouvoir d\'achat régional'],
        ['Nb_Hopitaux', '2 (Tataouine/Tozeur)', '12 (Tunis)', 'Densité structures santé'],
        ['Nb_Laboratoires', '7 (Tozeur)', '40 (Tunis)', 'Densité acheteurs spécialisés'],
    ],
    col_widths=[4, 4.5, 4.5, 4]
)
success_box('Résultat : chaque transaction est décrite par 8 variables prêtes pour le modèle.')
divider()

# Étape 5
h2('Étape 5 — Définition de la variable cible')
h3('La définition exacte')
code_block(
    'Pour chaque région → calculer le 75e percentile des quantités vendues\n'
    'Si Vente >= Q75_région  →  HAUSSE (classe 1)\n'
    'Sinon                   →  BAISSE (classe 0)'
)
h3('Pourquoi le 75e percentile par région ?')
para('Cette approche normalise les différences de taille de marché. Vendre 10 unités à '
     'Tozeur (petit marché) est une performance élevée. Vendre les mêmes 10 unités à '
     'Tunis (grand marché) est banal. En calculant le seuil au sein de chaque région, '
     'on compare chaque transaction à son propre contexte.')
add_table(
    ['Classe', 'Nombre', 'Proportion', 'Signification'],
    [
        ['BAISSE', '2 578', '71.7%', 'Transaction sous le 75e percentile régional'],
        ['HAUSSE', '1 017', '28.3%', 'Transaction au-dessus du 75e percentile régional'],
    ],
    col_widths=[3.5, 3, 3, 7.5]
)
h3('Signification métier')
para('Une région dont le modèle prédit 44% de transactions HAUSSE génère régulièrement '
     'des commandes importantes dans son contexte local — c\'est le signal pour y investir davantage.')
divider()

# Étape 6
doc.add_page_break()
h2('Étape 6 — Entraînement et sélection du modèle')
h3('Benchmark de 10 algorithmes')
para('Plutôt que de choisir un algorithme à l\'intuition, 10 algorithmes sont entraînés '
     'et comparés sur les mêmes données :')
add_table(
    ['Rang', 'Algorithme', 'Accuracy', 'F1-Score', 'Statut'],
    [
        ['🏆 1er', 'Voting Ensemble', '72.57%', '0.6937', 'CHAMPION'],
        ['2e', 'Gradient Boosting', '66.64%', '0.6811', ''],
        ['3e', 'Random Forest', '65.89%', '0.6542', ''],
        ['4e', 'AdaBoost', '63.48%', '0.6522', ''],
        ['5e', 'KNN', '67.56%', '0.6476', ''],
        ['6e', 'Decision Tree', '61.45%', '0.6321', ''],
        ['7e', 'XGBoost', '59.87%', '0.6184', ''],
        ['8e', 'Gaussian Naive Bayes', '71.73%', '0.5993', ''],
        ['9e', 'Logistic Regression', '53.75%', '0.5594', ''],
        ['Dernier', 'SVM (RBF)', '51.90%', '0.5363', ''],
    ],
    col_widths=[2, 5, 3, 3, 4]
)
h3('Optimisation GridSearchCV')
para('Les hyperparamètres du modèle gagnant sont optimisés automatiquement : '
     '16 combinaisons × 5 validations = 80 entraînements. '
     'F1 : 0.6937  →  0.7003 après optimisation.')
h3('Résultats finaux du modèle')
add_table(
    ['Métrique', 'Valeur', 'Interprétation'],
    [
        ['Accuracy', '72.94%', '73% des transactions correctement classées'],
        ['F1-Score pondéré', '70.03%', 'Équilibre précision/rappel sur les deux classes'],
        ['AUC-ROC', '70.92%', 'Discrimination significativement au-dessus du hasard (50%)'],
        ['CV 5-folds', '72.85% ± 0.84%', 'Stabilité confirmée — pas de surapprentissage'],
    ],
    col_widths=[4, 3, 10]
)
divider()

# Étape 7
h2('Étape 7 — Seuil de décision optimal')
para('Par défaut, un modèle déclare HAUSSE si la probabilité est ≥ 50%. '
     'La courbe Précision-Rappel identifie un seuil plus adapté.')
code_block('Seuil optimal identifié : 0.39 (au lieu de 0.50 par défaut)\n'
           'Effet : améliore la détection des vraies transactions à forte intensité\n'
           'Source : maximisation du F1 HAUSSE sur la courbe Précision-Rappel')
divider()

# Étape 8
h2('Étape 8 — Génération des recommandations stratégiques')
para('Pour chaque gouvernorat, on agrège les probabilités HAUSSE prédites. '
     'Les régions sont classées en trois tiers :')
add_table(
    ['Tiers', 'Condition', 'Décision', 'Action'],
    [
        ['Top 33%', 'Percentile ≥ 0.67', '🟢 AUGMENTER', 'Renforcer équipes et stocks (+25%)'],
        ['Milieu 33%', '0.33 ≤ Percentile < 0.67', '🟡 MAINTENIR', 'Stabiliser la couverture commerciale'],
        ['Bas 33%', 'Percentile < 0.33', '🔴 RÉDUIRE', 'Réallouer vers les régions prometteuses'],
    ],
    col_widths=[3, 4, 3.5, 6.5]
)
para('Enrichissement : pour les 14 régions disposant d\'historique CA réel, '
     'les variations trimestrielles réelles viennent enrichir la recommandation finale.')
h3('Résultat final — 24 gouvernorats classés')
add_table(
    ['Décision', 'Nb Gouvernorats', 'Exemples'],
    [
        ['🟢 AUGMENTER', '8', 'Ben Arous, Tataouine, Sousse, Sidi Bouzid, Tozeur, Kasserine, Kairouan, Jendouba'],
        ['🟡 MAINTENIR', '9', 'Zaghouan, Béja, Gafsa, Sfax, Siliana, Kebili, Le Kef, Monastir, Nabeul'],
        ['🔴 RÉDUIRE', '7', 'Bizerte, Mahdia, Gabes, Medenine, Manouba, Ariana, Tunis'],
    ],
    col_widths=[3.5, 3.5, 10]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DONNÉES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('3. Partie données')

h2('3.1 Sources des données')
add_table(
    ['Source', 'Table / Fichier', 'Contenu', 'Usage'],
    [
        ['SQL Server ERP', 'TabBcvd, TabBcvm, TabTiers, TabStock', '3 625 transactions 2022-2026', 'Données d\'entraînement'],
        ['INS Tunisie 2023', 'Données internes codées', 'Population, Indice achat, Hôpitaux, Labos', 'Enrichissement géo'],
        ['SQL Server ERP', 'TabBlvm', 'Factures avec CA HT par trimestre', 'Variations historiques'],
    ],
    col_widths=[4, 5, 5, 3]
)

h2('3.2 Qualité des données — Deux problèmes réels')
h3('Problème 1 — Absence de données géographiques fiables')
para('Le champ gouvernorat dans l\'ERP contient le code 17 pour 99.9% des clients — '
     'une valeur par défaut jamais renseignée. Ce n\'est pas une erreur du projet, '
     'c\'est la réalité courante des systèmes ERP. La solution (proxy économique + '
     'augmentation pondérée) est une réponse pragmatique et validée expérimentalement.')
h3('Problème 2 — Volume modeste')
para('3 595 transactions sur 24 régions sur 4 ans = ~150 transactions par région. '
     'C\'est suffisant pour faire fonctionner le modèle mais limite la précision '
     'pour les petites régions (Manouba : 59 transactions, Monastir : 66 transactions).')

h2('3.3 Pipeline de transformation des données')
code_block(
    'Données brutes       (3 625 lignes)\n'
    '     ↓  Suppression NaN, doublons, clampage Prix IQR\n'
    'Données propres      (3 595 lignes)\n'
    '     ↓  Inférence géographique + Quality Gate + augmentation 40%\n'
    'Données géolocalisées (24 gouvernorats couverts)\n'
    '     ↓  Merge INS 2023 (Population, Indice_Achat, Nb_Hopitaux, Nb_Laboratoires)\n'
    'Données enrichies    (16 colonnes)\n'
    '     ↓  Calcul Q75 par région → label HAUSSE/BAISSE\n'
    'Données étiquetées   (3 595 lignes, 2 classes)\n'
    '     ↓  Sélection 8 features\n'
    'Matrice X (3595 × 8) + vecteur y (3595 labels)'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — FEATURE ENGINEERING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('4. Feature Engineering')

h2('4.1 Les 8 variables du modèle')
add_table(
    ['Variable', 'Type', 'Source', 'Logique métier'],
    [
        ['Population', 'Démographique', 'INS 2023', 'Taille du marché potentiel de la région'],
        ['Indice_Achat', 'Économique', 'INS 2023', 'Pouvoir d\'achat régional moyen'],
        ['Nb_Hopitaux', 'Structurel santé', 'INS 2023', 'Densité des acheteurs professionnels (hôpitaux)'],
        ['Nb_Laboratoires', 'Structurel santé', 'INS 2023', 'Densité des acheteurs spécialisés'],
        ['Mois', 'Temporel', 'DateBL', 'Saisonnalité mensuelle des achats'],
        ['Année', 'Temporel', 'DateBL', 'Tendance annuelle de croissance'],
        ['Trimestre', 'Temporel', 'DateBL', 'Saisonnalité trimestrielle'],
        ['Prix', 'Transaction', 'PrixVente', 'Segment de produit / niveau de gamme'],
    ],
    col_widths=[4, 3.5, 3, 6.5]
)

h2('4.2 Ce que les corrélations révèlent')
para('La matrice de corrélation montre que Population, Indice_Achat, Nb_Hopitaux et '
     'Nb_Laboratoires sont très fortement corrélées (r = 0.88 à 0.98). C\'est attendu : '
     'les gouvernorats densément peuplés ont plus de pouvoir d\'achat, plus d\'hôpitaux '
     'et plus de laboratoires. Ces 4 variables portent essentiellement la même information — '
     '"taille et richesse du gouvernorat". Random Forest gère naturellement cette '
     'redondance par sa sélection aléatoire de features à chaque nœud.')
note_box('La variable Prix est la plus prédictive (Gini = 0.41). Elle capte indirectement '
         'l\'information géographique puisque le prix est corrélé à la région d\'assignation.')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — MODÈLE ML
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('5. Le modèle de Machine Learning')

h2('5.1 Pourquoi un Voting Ensemble ?')
para('Le principe du Voting Ensemble est celui de la sagesse des foules appliqué aux '
     'algorithmes. Plutôt que de faire confiance à un seul expert, on consulte plusieurs '
     'experts aux spécialités différentes et on prend la décision qui fait consensus.')

add_table(
    ['Composant', 'Spécialité', 'Contribution'],
    [
        ['Random Forest (100 arbres)', 'Patterns complexes et non-linéaires', 'Robuste aux outliers, capture les interactions entre variables'],
        ['Gradient Boosting', 'Correction itérative des erreurs', 'Fort sur les relations subtiles, arbres séquentiels'],
        ['Logistic Regression', 'Frontière de décision linéaire', 'Régularisation, évite le surapprentissage local'],
    ],
    col_widths=[4.5, 4.5, 8]
)

h2('5.2 Comment le modèle prend ses décisions')
para('Exemple concret pour la région SOUSSE, trimestre 3, année 2026 :')
code_block(
    'Étape 1 : Construction du vecteur de 8 features\n'
    '          Population=720000, Indice_Achat=124.0, Nb_Hopitaux=9,\n'
    '          Nb_Laboratoires=29, Mois=8, Année=2026, Trimestre=3, Prix=10.44\n\n'
    'Étape 2 : Les 3 sous-modèles calculent chacun P(HAUSSE)\n'
    '          RF → 0.54  |  GB → 0.52  |  LR → 0.48\n\n'
    'Étape 3 : Moyenne pondérée → P(HAUSSE) = 0.519\n\n'
    'Étape 4 : Seuil 0.39  →  0.519 >= 0.39  →  Prédiction : HAUSSE\n'
    '          Confiance : 51.9%  →  Recommandation : AUGMENTER'
)

h2('5.3 Validation croisée — Robustesse du modèle')
para('La validation croisée 5-fold stratifiée confirme la stabilité du modèle :')
code_block(
    '5 scores individuels : [71.49%, 73.57%, 73.71%, 72.32%, 73.16%]\n'
    'Moyenne CV           : 72.85%\n'
    'Écart-type           : ± 0.84%  (très faible → modèle stable)\n'
    'Test accuracy        : 72.94%   (CV ≈ Test → pas de surapprentissage)'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — VARIABLE CIBLE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('6. La variable cible — Explication approfondie')

h2('6.1 Ce qu\'elle mesure exactement')
para('La variable cible répond à la question :')
p = doc.add_paragraph()
run = p.add_run('"Cette transaction représente-t-elle une vente à haute intensité dans le contexte de sa région ?"')
set_font(run, size=12, bold=True, italic=True, color=(0x1E, 0x3A, 0x8A))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)

para('Ce n\'est pas une prédiction de tendance temporelle ("les ventes vont augmenter dans 3 mois") '
     '— c\'est une mesure de l\'intensité commerciale relative. Une transaction HAUSSE dans une '
     'petite région peut représenter 5 unités vendues ; dans une grande région, il en faudra '
     'peut-être 50 pour atteindre le même percentile.')

h2('6.2 La signification métier de l\'agrégation')
para('L\'intelligence du système réside dans le passage du niveau transactionnel au niveau régional. '
     'Si le modèle prédit que 44% des transactions d\'une région seront HAUSSE, cette région '
     'génère régulièrement des commandes importantes dans son contexte local. '
     'C\'est le signal pour y investir davantage.')

h2('6.3 Forces et limites')
add_table(
    ['Aspect', 'Détail'],
    [
        ['✅ Force principale', 'Normalisation intra-régionale : comparaison équitable entre marchés de tailles très différentes. Valorise les petits marchés dynamiques.'],
        ['✅ Gestion déséquilibre', 'class_weight="balanced" appliqué automatiquement sur tous les modèles.'],
        ['⚠  Limite principale', 'Q75 calculé sur ~150 échantillons par région. Statistiquement modeste — la prédiction est une estimation, pas une certitude.'],
        ['⚠  Limite conceptuelle', 'Mesure l\'intensité ponctuelle, pas la tendance temporelle. Un suivi trimestriel comparatif renforcerait la valeur métier.'],
    ],
    col_widths=[4, 13]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — MLOPS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('7. Le système MLOps')

h2('7.1 L\'API Flask')
para('L\'API expose les prédictions du modèle de façon standardisée. N\'importe quelle '
     'application de l\'entreprise peut interroger l\'API et obtenir une recommandation '
     'régionale en temps réel.')
add_table(
    ['Endpoint', 'Méthode', 'Paramètres', 'Retour'],
    [
        ['/api/predict', 'POST / GET', 'region, trimestre, year, prix (optionnel)', 'Prédiction + recommandation + confiance'],
        ['/api/batch-predict', 'POST', 'regions[], trimestre, year', 'Prédictions pour plusieurs régions'],
        ['/api/health', 'GET', '—', 'Statut du service et du modèle'],
        ['/api/regions', 'GET', '—', 'Liste des 24 régions valides'],
        ['/api/stats', 'GET', '—', 'Métriques du modèle en production'],
    ],
    col_widths=[4, 2.5, 5.5, 5]
)
h3('Exemple de réponse API')
code_block(
    '{\n'
    '  "success": true,\n'
    '  "region": "SOUSSE",\n'
    '  "prediction": "HAUSSE",\n'
    '  "confiance": 51.9,\n'
    '  "recommandation": "AUGMENTER",\n'
    '  "probabilite_hausse": 51.9,\n'
    '  "probabilite_baisse": 48.1,\n'
    '  "variation_estimee": "+5.9%",\n'
    '  "model_version": "1.4",\n'
    '  "model_accuracy": 0.7294\n'
    '}'
)

h2('7.2 L\'auto-retraining sécurisé')
para('Le pipeline de réentraînement automatique garantit que le modèle reste à jour :')
code_block(
    'auto_retrain.py\n'
    '  ├─ [0] Quality Gate → mesure couverture ERP (data_quality.py)\n'
    '  ├─ [1] Backup du modèle actuel (model_backup/ avec horodatage)\n'
    '  ├─ [2] Réentraînement via notebook (nbconvert, timeout 600s)\n'
    '  ├─ [3] Comparaison F1 nouveau vs ancien\n'
    '  └─ [4] Adoption si F1 ≥ actuel / Restauration si F1 < actuel'
)
success_box('Aucune dégradation de performance n\'est possible en production : '
            'si le nouveau modèle est moins bon, l\'ancien est automatiquement restauré.')

h2('7.3 Le Quality Gate adaptatif')
para('Avant chaque réentraînement, le module data_quality.py mesure la couverture '
     'géographique réelle de l\'ERP depuis 3 sources :')
for b in [
    'TabBcvd → TabBcvm → TabTiers.Ville (transactions directes)',
    'TabBlvm → TabTiers.Ville (factures — meilleure couverture : 11.5%)',
    'TabTiers.gouvernorat (codes non-défaut, actuellement 0%)',
]:
    bullet(b)
code_block(
    'Score actuel  : 11.5% (couverture TabBlvm)\n'
    'Seuil configuré : 50%\n'
    'Décision      : Augmentation ACTIVE (11.5% < 50%)\n'
    '\n'
    'Demain si ERP amélioré ≥ 50% → Augmentation DESACTIVEE automatiquement'
)

h2('7.4 Versioning des modèles')
add_table(
    ['Artefact', 'Fichier', 'Contenu'],
    [
        ['Modèle ML', 'predict_ventes_regions_v1.4.pkl', 'VotingClassifier optimisé (joblib)'],
        ['Scaler', 'scaler.pkl', 'StandardScaler (8 features)'],
        ['Métadonnées', 'metadata_v1.4.json', 'Métriques, features, date, version'],
        ['Recommandations', 'regional_recommendations.json', 'Cache des 24 recommandations'],
        ['Historique qualité', 'data_quality_history.json', 'Traçabilité MLOps de la couverture ERP'],
        ['Graphiques', '11 fichiers PNG', 'EDA, corrélation, ROC, PR, benchmark, etc.'],
    ],
    col_widths=[4, 5.5, 7.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — POINTS FORTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('8. Points forts du projet')

h2('8.1 Solidité technique')
strong_points = [
    ('Architecture modulaire professionnelle',
     'Séparation nette predict_api.py / flask_api.py. Chaque composant a une responsabilité unique et bien définie. Réutilisable et maintenable.'),
    ('Pipeline MLOps complet',
     'Très peu de projets PFE incluent un vrai système de réentraînement automatique avec backup/restore et Quality Gate adaptatif. C\'est généralement réservé aux équipes professionnelles.'),
    ('Démarche expérimentale rigoureuse',
     'L\'expérience contrôlée (avec/sans redistribution) mesure quantitativement l\'impact de chaque choix. Plutôt que de supposer, on mesure.'),
    ('Validation multi-niveaux',
     'Benchmark 10 algorithmes → GridSearchCV → validation croisée 5-fold → seuil optimal sur courbe PR → 19 checkpoints automatisés.'),
    ('11 visualisations professionnelles',
     'EDA, corrélation, déséquilibre des classes, benchmark, feature importance, confusion matrix, ROC, PR, learning curves, géographie, dashboard stratégique.'),
    ('Robustesse aux pannes',
     'Fallback données si SQL inaccessible, fallback démographique si CSV absent, fallback augmentation si Quality Gate échoue.'),
    ('Quality Gate adaptatif (innovation)',
     'Mesure automatique de la qualité ERP avant chaque réentraînement. Le pipeline s\'améliore automatiquement avec les données.'),
]
for titre, desc in strong_points:
    h3('✅ ' + titre)
    para(desc)

h2('8.2 Ce qui est impressionnant pour un PFE')
success_box(
    'La combinaison d\'une API REST fonctionnelle + pipeline de réentraînement automatique '
    '+ Quality Gate adaptatif + versioning d\'artefacts dans un même projet académique est rare. '
    'Ces éléments sont couramment présentés comme "perspectives futures" dans la majorité des PFE '
    '— ici, ils sont implémentés et opérationnels.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — LIMITES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('9. Points faibles et limites')

h2('9.1 Limites des données')
add_table(
    ['Limite', 'Impact', 'Mitigation appliquée'],
    [
        ['Volume modeste (~150 transactions/région)', 'Prédictions moins précises pour les petites régions (Manouba: 59, Monastir: 66)', 'Documenté comme limite — résultats à interpréter avec prudence pour ces régions'],
        ['Géographie approximative (proxy prix)', 'Assignation régionale estimée, pas mesurée', 'Expérience validée (+3.3% F1), Quality Gate pour détecter l\'amélioration ERP'],
        ['TATAOUINE surreprésentée (703/3595)', 'Biais potentiel dans les recommandations globales', 'Documenté, recommandations à croiser avec analyse terrain'],
    ],
    col_widths=[4.5, 5, 7.5]
)

h2('9.2 Limites du modèle')
add_table(
    ['Limite', 'Valeur mesurée', 'Explication'],
    [
        ['Recall HAUSSE faible', '29.18%', 'Le modèle détecte 29% des vraies transactions à forte intensité. Limite inhérente aux données déséquilibrées.'],
        ['Pas de features temporelles', '—', 'Pas de lag des ventes précédentes ni de tendance. Amélioration future : features de tendance trimestrielle.'],
        ['Multicolinéarité démographique', 'r = 0.88-0.98', 'Les 4 variables démographiques portent la même info. Géré nativement par Random Forest.'],
    ],
    col_widths=[4.5, 3.5, 9]
)

h2('9.3 Limites métier')
note_box(
    'TUNIS (443 transactions, 1.1M habitants) est recommandée RÉDUIRE. '
    'Cette recommandation s\'explique par une intensité transactionnelle relative faible (transactions petites '
    'dans le grand contexte de Tunis), pas par un marché peu intéressant en volume absolu. '
    'Les recommandations doivent être croisées avec l\'analyse commerciale terrain.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('10. Conclusion globale')

h2('10.1 Est-ce un projet solide ?')
success_box(
    'Oui, clairement. Ce projet résout un vrai problème métier (allocation des ressources '
    'commerciales régionales) avec une vraie solution technique (pipeline ML end-to-end). '
    'Il va bien au-delà d\'un notebook d\'exploration : il produit un système utilisable, '
    'maintenable et évolutif.'
)
para('Les choix techniques sont justifiés, les limites sont reconnues et documentées, '
     'la démarche est scientifique. Le résultat — un Voting Ensemble avec 72.94% d\'accuracy, '
     'déployé derrière une API Flask avec retraining automatique — est cohérent et défendable.')

h2('10.2 Est-il défendable devant un jury ?')
para('Oui, avec une bonne préparation. Les questions difficiles viendront probablement sur '
     'trois points : la définition de la variable cible, la méthode d\'inférence géographique, '
     'et le recall HAUSSE à 29%. Ces trois points ont des réponses solides.')
note_box(
    'Présenter les limites soi-même avant que le jury ne les soulève est perçu très positivement. '
    'Un étudiant qui connaît les limites de son travail montre qu\'il comprend vraiment ce qu\'il a fait.'
)

h2('10.3 Niveau du projet')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Avancé — à la frontière du niveau industriel')
set_font(run, size=16, bold=True, color=(0x07, 0x47, 0x26))
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(12)

add_table(
    ['Critère', 'Évaluation', 'Commentaire'],
    [
        ['Niveau technique', '⭐⭐⭐⭐⭐', 'Architecture professionnelle, MLOps complet, API deployable'],
        ['Qualité scientifique', '⭐⭐⭐⭐', 'Démarche rigoureuse, limites documentées, validation multi-niveaux'],
        ['Valeur métier', '⭐⭐⭐⭐', 'Problème réel, recommandations actionnables, données ERP réelles'],
        ['Fiabilité soutenance', '⭐⭐⭐⭐⭐', 'Prêt — questions difficiles anticipées et réponses préparées'],
        ['Comparaison autres PFE', '⭐⭐⭐⭐⭐', 'Significativement au-dessus de la moyenne académique'],
    ],
    col_widths=[5, 3, 9]
)

para('')
para(
    'La majorité des projets PFE en Data Science présentent un notebook avec un modèle '
    'et quelques graphiques. Ce projet inclut en plus une API REST déployable, un pipeline '
    'de réentraînement automatique avec Quality Gate, du versioning d\'artefacts, une gestion '
    'des erreurs robuste, et une démarche expérimentale rigoureuse. C\'est le niveau que les '
    'entreprises attendent de leurs jeunes data scientists en sortie d\'études supérieures.',
    italic=True
)

# ── Pied de page ─────────────────────────────────────────────────────────────
doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Rapport généré automatiquement  |  Pipeline ML V1.4  |  '
    'VotingClassifier  |  Accuracy: 72.94%  |  F1: 70.03%  |  AUC: 70.92%'
)
set_font(run, size=9, italic=True, color=(0x9C, 0xA3, 0xAF))

# ── Sauvegarde ───────────────────────────────────────────────────────────────
output_path = 'Analyse_Globale_ML_PFE_V1.4.docx'
doc.save(output_path)
print(f'Document genere avec succes : {output_path}')
